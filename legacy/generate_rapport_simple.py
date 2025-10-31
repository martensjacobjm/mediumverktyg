#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Genererar Word-rapport UTAN automatiska diagram (diagram läggs till manuellt)
Denna version kräver INTE CoolProp eller matplotlib
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

# Hitta output-mapp
output_dir = os.path.join(os.path.dirname(__file__), 'outputs')
os.makedirs(output_dir, exist_ok=True)

# Skapa dokument
doc = Document()

# Definiera färger
COLOR_BLUE = RGBColor(0, 51, 102)
COLOR_GRAY = RGBColor(89, 89, 89)

print("\n" + "="*70)
print("GENERERAR WORD-RAPPORT (UTAN AUTOMATISKA DIAGRAM)")
print("="*70)

# ============================================================================
# FRAMSIDA
# ============================================================================

# Titel
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('ARBETSMEDIUM-ANALYS\n')
title_run.font.size = Pt(26)
title_run.font.bold = True
title_run.font.color.rgb = COLOR_BLUE

subtitle_run = title.add_run('Tesla-Turbin ORC-System för Lågtemperaturapplikation')
subtitle_run.font.size = Pt(16)
subtitle_run.font.color.rgb = COLOR_GRAY

doc.add_paragraph()  # Mellanrum

# Projektinfo
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_text = [
    'ORC Malung',
    'Temperaturområde: 30-80°C',
    'Effektmål: 1-2 kW elektrisk',
    '',
    f'Datum: {datetime.now().strftime("%Y-%m-%d")}',
    'Version: 2.0 - Rapport med manuella diagram'
]
for line in info_text:
    run = info.add_run(line + '\n')
    run.font.size = Pt(12)

doc.add_page_break()

# ============================================================================
# EXECUTIVE SUMMARY
# ============================================================================

doc.add_heading('Executive Summary', 0)

summary_text = [
    ('Syfte och Omfattning',
     'Detta dokument presenterar en systematisk termodynamisk analys för val av arbetsmedium '
     'till ett Tesla-turbin baserat ORC-system (Organic Rankine Cycle). Systemet är designat '
     'för lågtemperaturapplikationer (30-80°C) med målsättning att generera 1-2 kW elektrisk '
     'effekt från värmepump, solfångare och/eller vedeldning.'),

    ('Metodologi',
     'Analysen baseras på termodynamiska beräkningar med CoolProp-databas, säkerhetsbedömning '
     'enligt ASHRAE Standard 34-2019, miljöanalys enligt EU F-gas Regulation 517/2014, samt '
     'viskositetsanpassning för Tesla-turbingeometri. Tre primära kandidater har utvärderats: '
     'R1233zd(E), R245fa och R1234ze(Z).'),

    ('Huvudresultat',
     'R1233zd(E) rekommenderas som primärt arbetsmedium baserat på optimal kokpunkt (19,0°C), '
     'lägst drifttryck (2,93 bar vid 50°C), säkraste klassning (A1), och nästan noll '
     'klimatpåverkan (GWP <7). För 1 kW eleffekt vid drift 50°C → 20°C krävs 9,1 g/s massflöde, '
     '1,96 kW förångare och optimalt diskavstånd 0,191 mm.'),

    ('Rekommendation',
     'Implementera R1233zd(E) som primärt medium med R245fa som backup. Merkostnad cirka 200-400 € '
     'motiveras av överlägsen säkerhet (A1 vs B1), 147× lägre klimatpåverkan (GWP <7 vs 1030), '
     'och 15% lägre drifttryck som förenklar systemdesign.')
]

for heading, text in summary_text:
    p = doc.add_paragraph()
    run = p.add_run(heading + ': ')
    run.font.bold = True
    run.font.size = Pt(11)
    p.add_run(text).font.size = Pt(11)

doc.add_page_break()

# ============================================================================
# INNEHÅLLSFÖRTECKNING
# ============================================================================

doc.add_heading('Innehållsförteckning', 0)

toc_items = [
    '1. Inledning och Bakgrund',
    '2. Problemställning',
    '3. Systemkrav och Driftförhållanden',
    '4. Metodologi',
    '5. Kandidat-medier',
    '6. Termodynamisk Analys',
    '7. Viskositet och Tesla-Turbin Anpassning',
    '8. Säkerhet och Miljö',
    '9. Dimensionering och Beräkningar',
    '10. Diskussion och Jämförelse',
    '11. Slutsatser och Rekommendation',
    '12. Bilagor: Diagram och Tabeller',
    '13. Referenser'
]

for item in toc_items:
    p = doc.add_paragraph(item, style='List Number')
    p.paragraph_format.left_indent = Inches(0.5)

doc.add_page_break()

# ============================================================================
# KAPITEL 1-5
# ============================================================================

doc.add_heading('1. Inledning och Bakgrund', 1)
p = doc.add_paragraph(
    'ORC (Organic Rankine Cycle) är en etablerad teknologi för konvertering av lågtemperaturvärme '
    'till elektrisk energi. ORC Malung-projektet utvecklar ett småskaligt system med Tesla-turbin, '
    'multipla värmekällor (värmepump 30-60°C, solfångare 40-80°C, ved/pellets 60-80°C), och '
    'måleffekt 1-2 kW elektrisk.'
)

doc.add_heading('2. Problemställning', 1)
p = doc.add_paragraph(
    'Arbetsmediet påverkar systemtryck, turbingeometri, verkningsgrad, säkerhet, och miljöpåverkan. '
    'Lågtemperatursystem (30-80°C) kräver optimal kokpunkt nära kondenseringszonen (10-30°C), '
    'lämplig viskositet för Tesla-turbin diskavstånd, och acceptabel säkerhet för heminstallation.'
)

doc.add_heading('3. Systemkrav och Driftförhållanden', 1)

# Tabell
table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'
header_cells = table.rows[0].cells
header_cells[0].text = 'Komponent'
header_cells[1].text = 'Temperatur'
header_cells[2].text = 'Tryck (R1233zd(E))'
data = [
    ['Förångare', '50°C (30-80°C)', '2,93 bar'],
    ['Kondensor vinter', '20°C', '1,08 bar'],
    ['Kondensor sommar KB', '10°C', '0,73 bar']
]
for i, row_data in enumerate(data, 1):
    cells = table.rows[i].cells
    for j, value in enumerate(row_data):
        cells[j].text = value

doc.add_paragraph()

doc.add_heading('4. Metodologi', 1)
p = doc.add_paragraph(
    'Termodynamiska beräkningar med CoolProp 7.1.0 (validerad mot NIST REFPROP), '
    'säkerhetsbedömning enligt ASHRAE Standard 34-2019, miljöanalys enligt EU F-gas, '
    'och diskavståndsberäkning med gränsskiktsteori.'
)

doc.add_heading('5. Kandidat-medier', 1)

candidates_text = [
    ('R1233zd(E)', 'Kokpunkt 19,0°C, ASHRAE A1, GWP <7. Modern lågtemp-ORC favorit.'),
    ('R245fa', 'Kokpunkt 15,3°C, ASHRAE B1, GWP 1030. Etablerad sedan 20 år.'),
    ('R1234ze(Z)', 'Kokpunkt 9,8°C, ASHRAE A2L, GWP <1. Kräver säkerhetsanalys.')
]

for name, desc in candidates_text:
    p = doc.add_paragraph()
    run = p.add_run(name + ': ')
    run.font.bold = True
    p.add_run(desc)

doc.add_page_break()

# ============================================================================
# KAPITEL 6 - Med platshållare för diagram
# ============================================================================

doc.add_heading('6. Termodynamisk Analys', 1)

doc.add_heading('6.1 Tryck-Temperatur Översikt', 2)

p = doc.add_paragraph(
    'Figur 6.1 visar tryck-temperatur kurvor för de två huvudkandidaterna. R1233zd(E) har lägst tryck '
    'vid alla temperaturer, vilket förenklar systemdesign och minskar komponentkostnader.'
)

# Platshållare för diagram
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('[INFOGA DIAGRAM HÄR: Tryck-Temperatur jämförelse]')
run.font.bold = True
run.font.color.rgb = RGBColor(255, 0, 0)
run.font.size = Pt(12)

caption = doc.add_paragraph()
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = caption.add_run('Figur 6.1: Tryck-Temperatur jämförelse för R245fa och R1233zd(E)')
run.font.size = Pt(10)
run.font.italic = True

doc.add_paragraph()

doc.add_heading('6.2 Detaljerad Egenskapsjämförelse', 2)

p = doc.add_paragraph(
    'Figur 6.2 visar fyra kritiska parametrar: mättningstryck, förångningsvärme (hfg), '
    'viskositet (påverkar diskavstånd), och ångdensitet. Alla parametrar är viktiga för '
    'systemprestanda och dimensionering.'
)

# Platshållare för diagram 2
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('[INFOGA DIAGRAM HÄR: 4-panel termodynamisk jämförelse]')
run.font.bold = True
run.font.color.rgb = RGBColor(255, 0, 0)
run.font.size = Pt(12)

caption = doc.add_paragraph()
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = caption.add_run('Figur 6.2: Termodynamisk jämförelse - Tryck, hfg, Viskositet, Densitet')
run.font.size = Pt(10)
run.font.italic = True

doc.add_paragraph()

# Nyckeldata tabell
doc.add_heading('6.3 Nyckeldata vid Drifttemperaturer', 2)

comparison_table = doc.add_table(rows=7, cols=5)
comparison_table.style = 'Medium Grid 1 Accent 1'

headers = comparison_table.rows[0].cells
headers[0].text = 'Temperatur'
headers[1].text = 'Medium'
headers[2].text = 'Tryck [bar]'
headers[3].text = 'hfg [kJ/kg]'
headers[4].text = 'μ_ånga [μPa·s]'

comp_data = [
    ('10°C', 'R1233zd(E)', '0,73', '198,8', '10,5'),
    ('', 'R245fa', '0,82', '199,5', '11,2'),
    ('50°C', 'R1233zd(E)', '2,93', '177,3', '12,1'),
    ('', 'R245fa', '3,44', '175,9', '12,9'),
    ('80°C', 'R1233zd(E)', '6,58', '157,9', '13,4'),
    ('', 'R245fa', '7,89', '153,9', '14,3'),
]

for i, row_data in enumerate(comp_data, 1):
    cells = comparison_table.rows[i].cells
    for j, value in enumerate(row_data):
        cells[j].text = value

doc.add_page_break()

# ============================================================================
# RESTEN AV KAPITLEN (samma som tidigare)
# ============================================================================

doc.add_heading('7. Viskositet och Tesla-Turbin Anpassning', 1)

p = doc.add_paragraph(
    'Optimalt diskavstånd för Tesla-turbin bestäms av mediets viskositet enligt gränsskiktsteori. '
    'R1233zd(E) och R245fa har mycket liknande viskositeter (12,1 vs 12,9 μPa·s vid 50°C).'
)

# Diskavstånd tabell
disc_table = doc.add_table(rows=4, cols=4)
disc_table.style = 'Light Grid Accent 1'

disc_header = disc_table.rows[0].cells
disc_header[0].text = 'Medium'
disc_header[1].text = 'μ vid 50°C [μPa·s]'
disc_header[2].text = 'Skalningsfaktor'
disc_header[3].text = 'Diskavstånd [mm]'

disc_data = [
    ['Luft (TesTur ref)', '18,2', '1,000', '0,234'],
    ['R1233zd(E)', '12,1', '0,816', '0,191'],
    ['R245fa', '12,9', '0,842', '0,197']
]

for i, row_data in enumerate(disc_data, 1):
    cells = disc_table.rows[i].cells
    for j, value in enumerate(row_data):
        cells[j].text = value

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('Slutsats: ')
run.font.bold = True
p.add_run(
    'Praktiskt identiska diskavstånd (0,19-0,20 mm) innebär att samma turbindesign '
    'kan användas för båda medier.'
)

doc.add_page_break()

# ============================================================================
# SÄKERHET
# ============================================================================

doc.add_heading('8. Säkerhet och Miljö', 1)

doc.add_heading('8.1 ASHRAE Säkerhetsklassning', 2)

safety_table = doc.add_table(rows=4, cols=4)
safety_table.style = 'Light Grid Accent 1'

safety_header = safety_table.rows[0].cells
safety_header[0].text = 'Medium'
safety_header[1].text = 'ASHRAE Klass'
safety_header[2].text = 'Toxicitet'
safety_header[3].text = 'Brandfarlighet'

safety_data = [
    ['R1233zd(E)', 'A1', 'Mycket låg', 'Ej brandfarlig'],
    ['R245fa', 'B1', 'Låg', 'Ej brandfarlig'],
    ['R1234ze(Z)', 'A2L', 'Mycket låg', 'Lätt brandfarlig']
]

for i, row_data in enumerate(safety_data, 1):
    cells = safety_table.rows[i].cells
    for j, value in enumerate(row_data):
        cells[j].text = value

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('R1233zd(E): A1 - ')
run.font.bold = True
p.add_run('Säkrast möjliga klassning. Inga speciella säkerhetsåtgärder utöver standard F-gas krav.')

doc.add_heading('8.2 Miljöpåverkan', 2)

env_table = doc.add_table(rows=4, cols=3)
env_table.style = 'Light Grid Accent 1'

env_header = env_table.rows[0].cells
env_header[0].text = 'Medium'
env_header[1].text = 'GWP (100 år)'
env_header[2].text = 'Framtidssäker'

env_data = [
    ['R1233zd(E)', '<7', 'Ja (mot 2030 regler)'],
    ['R245fa', '1030', 'Osäkert'],
    ['R1234ze(Z)', '<1', 'Ja']
]

for i, row_data in enumerate(env_data, 1):
    cells = env_table.rows[i].cells
    for j, value in enumerate(row_data):
        cells[j].text = value

doc.add_paragraph()

p = doc.add_paragraph(
    'R1233zd(E) har 147× lägre klimatpåverkan än R245fa och är framtidssäker mot '
    'kommande strängare F-gas regleringar (EU mål: GWP <150 till 2030).'
)

doc.add_page_break()

# ============================================================================
# DIMENSIONERING
# ============================================================================

doc.add_heading('9. Dimensionering och Beräkningar', 1)

doc.add_heading('9.1 Grunddimensionering (R1233zd(E) 50°C → 20°C, 1 kW)', 2)

dim_table = doc.add_table(rows=9, cols=3)
dim_table.style = 'Medium Grid 1 Accent 1'

dim_header = dim_table.rows[0].cells
dim_header[0].text = 'Parameter'
dim_header[1].text = 'Värde'
dim_header[2].text = 'Kommentar'

dim_data = [
    ['Måleffekt (el)', '1,0 kW', 'Efter generatorförluster'],
    ['Massflöde', '9,1 g/s', 'Från hfg och η_turb=55%'],
    ['Förångare', '1,96 kW', 'Värmeutvinning från tank'],
    ['Kondensor', '2,96 kW', 'Värmebortförsel'],
    ['Tryckförhållande', '2,71:1', 'Optimal för Tesla-turbin'],
    ['Pumpeffekt', '2,0 W', 'Försumbar (0,2%)'],
    ['Diskavstånd', '0,191 mm', 'Skalat från TesTur'],
    ['Köldbärare', '8,5 L/min', 'Sommardrift ΔT=5K']
]

for i, row_data in enumerate(dim_data, 1):
    cells = dim_table.rows[i].cells
    for j, value in enumerate(row_data):
        cells[j].text = value

doc.add_paragraph()

doc.add_heading('9.2 Skalning och Maximal Prestanda', 2)

p = doc.add_paragraph()
run = p.add_run('Skalning till 2 kW: ')
run.font.bold = True
p.add_run('Massflöde 18,2 g/s, Förångare 3,91 kW, Kondensor 5,91 kW, Köldbärare 17 L/min.')

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('Maximal prestanda (80°C → 10°C sommardrift): ')
run.font.bold = True
p.add_run(
    'Tryckförhållande 8,96:1, Carnot 19,8% (dubbelt), Systemverkningsgrad 8-10% (dubbelt). '
    'OPTIMAL konfiguration för högsta elproduktion.'
)

doc.add_page_break()

# ============================================================================
# DISKUSSION
# ============================================================================

doc.add_heading('10. Diskussion och Jämförelse', 1)

doc.add_heading('10.1 R1233zd(E) vs R245fa', 2)

compare_table = doc.add_table(rows=7, cols=3)
compare_table.style = 'Medium Grid 1 Accent 1'

comp_header = compare_table.rows[0].cells
comp_header[0].text = 'Parameter'
comp_header[1].text = 'R1233zd(E)'
comp_header[2].text = 'R245fa'

comp_data2 = [
    ['Kokpunkt', '19,0°C (BÄTTRE)', '15,3°C'],
    ['Tryck 50°C', '2,93 bar (BÄTTRE)', '3,44 bar (+17%)'],
    ['Massflöde 1 kW', '9,1 g/s', '9,0 g/s (samma)'],
    ['ASHRAE', 'A1 (BÄTTRE)', 'B1'],
    ['GWP', '<7 (MYCKET BÄTTRE)', '1030 (147× högre)'],
    ['Kostnad', '+200-400 € (+20%)', 'Referens']
]

for i, row_data in enumerate(comp_data2, 1):
    cells = compare_table.rows[i].cells
    for j, value in enumerate(row_data):
        cells[j].text = value

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('Analys: ')
run.font.bold = True
p.add_run(
    'Merkostnad 200-400 € (2-4% av totalkostnad 10-15 k€) motiveras av 15% lägre tryck, '
    'bättre säkerhet (A1 vs B1), och 147× lägre klimatpåverkan.'
)

doc.add_heading('10.2 Varför INTE R1234ze(Z)?', 2)

p = doc.add_paragraph(
    'Trots lägst GWP (<1) rekommenderas INTE för första prototyp: A2L kräver läckdetektorer '
    'och ventilation, högre tryck (5,62 bar), svårare kondensering (tryck 2,80 bar vid 20°C), '
    'och sämre tillgänglighet.'
)

doc.add_page_break()

# ============================================================================
# SLUTSATSER
# ============================================================================

doc.add_heading('11. Slutsatser och Rekommendation', 1)

doc.add_heading('11.1 Primär Rekommendation: R1233zd(E)', 2)

p = doc.add_paragraph()
run = p.add_run('R1233zd(E)')
run.font.bold = True
run.font.size = Pt(12)
p.add_run(' rekommenderas som primärt arbetsmedium baserat på:')

motivations = [
    'Optimal kokpunkt 19,0°C (närmare kondensering 10-30°C)',
    'Lägst drifttryck 2,93 bar vid 50°C (15% bättre än R245fa)',
    'Säkrast klassning A1 (lägst toxicitet, ej brandfarlig)',
    'Nästan noll klimatpåverkan GWP <7 (147× bättre än R245fa)',
    'Framtidssäker mot kommande F-gas regleringar',
    'Merkostnad 200-400 € försumbar (2-4% av totalkostnad)'
]

for item in motivations:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('11.2 Tekniska Specifikationer', 2)

spec_table = doc.add_table(rows=7, cols=2)
spec_table.style = 'Medium Shading 1 Accent 1'

spec_data = [
    ['Arbetsmedium', 'R1233zd(E) (primär) / R245fa (backup)'],
    ['Massflöde 1 kW', '9-10 g/s'],
    ['Förångare', '2 kW, 2,9 bar designtryck'],
    ['Kondensor', '3 kW, 1,1 bar designtryck'],
    ['Diskavstånd', '0,19-0,20 mm (startpunkt)'],
    ['Köldbärare sommar', '8-10 L/min vid ΔT=5K']
]

for row_data in spec_data:
    cells = spec_table.add_row().cells
    cells[0].text = row_data[0]
    cells[1].text = row_data[1]

doc.add_page_break()

# ============================================================================
# REFERENSER
# ============================================================================

doc.add_heading('13. Referenser', 1)

doc.add_heading('13.1 Termodynamiska Databaser', 2)

refs_thermo = [
    'Bell, I.H., et al. (2014). "Pure and Pseudo-pure Fluid Thermophysical Property Evaluation '
    'and CoolProp". Ind. Eng. Chem. Res., 53(6):2498-2508.',

    'NIST REFPROP Database Version 10.0, National Institute of Standards and Technology.'
]

for ref in refs_thermo:
    p = doc.add_paragraph(ref)
    p.paragraph_format.left_indent = Inches(0.5)

doc.add_heading('13.2 Tesla-Turbin Forskning', 2)

refs_tesla = [
    'Lampart, P., et al. (2019). "Design Analysis of Tesla Micro-Turbine". Energies, 12(44).',

    'Reiley, Ken (2010-2020). Tesla Turbine Experimental Research.',

    'Tesla, N. (1913). "Turbine." British Patent GB179043.'
]

for ref in refs_tesla:
    p = doc.add_paragraph(ref)
    p.paragraph_format.left_indent = Inches(0.5)

doc.add_heading('13.3 Standards', 2)

refs_standards = [
    'ASHRAE Standard 34-2019: Designation and Safety Classification of Refrigerants.',

    'EU Regulation 517/2014: F-gas Regulation.',

    'SS-EN 378:2016: Refrigerating systems - Safety requirements.'
]

for ref in refs_standards:
    p = doc.add_paragraph(ref)
    p.paragraph_format.left_indent = Inches(0.5)

# ============================================================================
# SPARA
# ============================================================================

output_path = os.path.join(output_dir, 'ORC_Arbetsmedium_Analys.docx')
doc.save(output_path)

print(f"\n{'='*70}")
print("WORD-DOKUMENT SKAPAT!")
print('='*70)
print(f"\nFilnamn: ORC_Arbetsmedium_Analys.docx")
print(f"Full soekvaeg: {os.path.abspath(output_path)}")
print(f"\nOBS: Diagram-platshållare är röda - lägg till diagram manuellt!")
print(f"{'='*70}\n")
